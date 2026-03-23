"""
Safety Standards RAG Chatbot — Streamlit Community Cloud Edition
Stack: Mistral AI API · LangChain · ChromaDB (in-memory) · Streamlit

Document parsing:
  • pdfplumber   — text + structured tables (preserves rows/cols as markdown)
  • pytesseract  — OCR for scanned / image-only PDF pages
  • pdf2image    — rasterises PDF pages for OCR
  • python-docx  — DOCX text + proper table cell extraction
  • docx2txt     — fallback for legacy .doc files

"""

import os, time, logging, io, re
from pathlib import Path

import streamlit as st

st.set_page_config(
    page_title="SafetyRAG",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ── Constants ─────────────────────────────────────────────────────────────────
MISTRAL_MODELS = {
    "Mistral Small (fast, free tier)": "mistral-small-latest",
    "Mistral Medium":                  "mistral-medium-latest",
    "Mixtral 8x7B":                    "open-mixtral-8x7b",
    "Mistral 7B (open)":               "open-mistral-7b",
}
EMBED_MODEL  = "sentence-transformers/all-MiniLM-L6-v2"
ANSWER_MODES = {
    "Concise":       "Give a concise, direct answer in 2-4 sentences.",
    "Detailed":      "Give a thorough, detailed answer covering all relevant aspects.",
    "Bullet Points": "Structure your entire answer as clear bullet points.",
}
MEMORY_TURNS = 4

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Sans:wght@300;400;500;600&family=JetBrains+Mono:wght@400;500&display=swap');
:root {
    --bg:#0d0f14;--surface:#151820;--surface2:#1c2030;--border:#252a3a;
    --accent:#f0a500;--accent2:#e05c2a;--text:#e8eaf0;--muted:#6b7280;
    --success:#22c55e;--warn:#f59e0b;--danger:#ef4444;--user-bg:#1e2540;
}
html,body,[data-testid="stAppViewContainer"]{background:var(--bg)!important;color:var(--text)!important;font-family:'DM Sans',sans-serif!important;}
[data-testid="stSidebar"]{background:var(--surface)!important;border-right:1px solid var(--border)!important;}
.rag-header{display:flex;align-items:center;gap:14px;padding:24px 0 18px;border-bottom:1px solid var(--border);margin-bottom:20px;}
.rag-logo{width:46px;height:46px;flex-shrink:0;background:linear-gradient(135deg,var(--accent),var(--accent2));border-radius:12px;display:flex;align-items:center;justify-content:center;font-size:22px;box-shadow:0 4px 18px rgba(240,165,0,.28);}
.rag-title{font-family:'DM Serif Display',serif!important;font-size:26px!important;font-weight:400!important;color:var(--text)!important;margin:0!important;}
.rag-subtitle{font-size:11px;color:var(--muted);letter-spacing:.08em;text-transform:uppercase;}
.status-badge{display:inline-flex;align-items:center;gap:6px;padding:4px 12px;border-radius:20px;font-size:12px;font-weight:500;}
.status-ready{background:rgba(34,197,94,.12);color:var(--success);border:1px solid rgba(34,197,94,.25);}
.status-loading{background:rgba(240,165,0,.12);color:var(--accent);border:1px solid rgba(240,165,0,.25);}
.chat-wrapper{display:flex;flex-direction:column;gap:16px;padding:4px 0;}
.msg-user,.msg-bot{display:flex;gap:12px;align-items:flex-start;animation:fadeUp .3s ease;}
.msg-user{flex-direction:row-reverse;}
@keyframes fadeUp{from{opacity:0;transform:translateY(6px)}to{opacity:1;transform:translateY(0)}}
.msg-avatar{width:34px;height:34px;border-radius:9px;flex-shrink:0;display:flex;align-items:center;justify-content:center;font-size:15px;}
.avatar-user{background:var(--user-bg);border:1px solid var(--border);}
.avatar-bot{background:linear-gradient(135deg,var(--accent),var(--accent2));box-shadow:0 2px 10px rgba(240,165,0,.22);}
.msg-bubble{max-width:78%;padding:13px 17px;border-radius:13px;font-size:14px;line-height:1.7;}
.bubble-user{background:var(--user-bg);border:1px solid var(--border);border-top-right-radius:4px;}
.bubble-bot{background:var(--surface);border:1px solid var(--border);border-top-left-radius:4px;}
.msg-meta{display:flex;align-items:center;gap:8px;margin-top:6px;flex-wrap:wrap;}
.meta-time{font-size:11px;color:var(--muted);font-family:'JetBrains Mono',monospace;}
.conf-high{font-size:11px;padding:2px 8px;border-radius:10px;background:rgba(34,197,94,.12);color:var(--success);border:1px solid rgba(34,197,94,.25);}
.conf-med{font-size:11px;padding:2px 8px;border-radius:10px;background:rgba(245,158,11,.12);color:var(--warn);border:1px solid rgba(245,158,11,.25);}
.conf-low{font-size:11px;padding:2px 8px;border-radius:10px;background:rgba(239,68,68,.12);color:var(--danger);border:1px solid rgba(239,68,68,.25);}
.parse-badge{font-size:10px;padding:1px 6px;border-radius:8px;background:rgba(240,165,0,.08);border:1px solid rgba(240,165,0,.2);color:#c9a84c;margin:1px;font-family:'JetBrains Mono',monospace;}
.sources-block{margin-top:10px;padding:9px 13px;background:var(--surface2);border-left:3px solid var(--accent);border-radius:0 7px 7px 0;font-size:12px;color:var(--muted);}
.sources-block strong{color:var(--accent);}
.source-tag{display:inline-block;background:rgba(240,165,0,.08);border:1px solid rgba(240,165,0,.2);border-radius:4px;padding:2px 7px;margin:2px 2px 0 0;font-family:'JetBrains Mono',monospace;font-size:11px;color:#c9a84c;}
.chunk-box{background:var(--bg);border:1px solid var(--border);border-radius:7px;padding:10px 13px;font-size:11.5px;font-family:'JetBrains Mono',monospace;color:var(--muted);white-space:pre-wrap;max-height:180px;overflow-y:auto;margin-bottom:8px;}
.chunk-label{font-size:11px;font-weight:600;color:var(--accent);margin-bottom:4px;}
.parse-info{background:var(--surface2);border:1px solid var(--border);border-radius:8px;padding:8px 12px;font-size:11.5px;color:var(--muted);margin-top:6px;}
.stTextInput>div>div>input{background:var(--surface)!important;border:1px solid var(--border)!important;border-radius:11px!important;color:var(--text)!important;font-family:'DM Sans',sans-serif!important;font-size:14px!important;padding:13px 17px!important;transition:border-color .2s!important;}
.stTextInput>div>div>input:focus{border-color:var(--accent)!important;box-shadow:0 0 0 3px rgba(240,165,0,.1)!important;}
.stTextInput>div>div>input::placeholder{color:var(--muted)!important;}
.stButton>button{background:linear-gradient(135deg,var(--accent),var(--accent2))!important;color:#0d0f14!important;border:none!important;border-radius:9px!important;font-weight:600!important;font-family:'DM Sans',sans-serif!important;transition:opacity .2s,transform .15s!important;box-shadow:0 3px 12px rgba(240,165,0,.18)!important;}
.stButton>button:hover{opacity:.87!important;transform:translateY(-1px)!important;}
.sidebar-label{font-size:11px;font-weight:600;letter-spacing:.1em;text-transform:uppercase;color:var(--muted);margin-bottom:8px;display:block;}
.stat-row{display:flex;justify-content:space-between;margin-bottom:5px;}
.stat-key{font-size:12px;color:var(--muted);}
.stat-val{font-size:12px;font-weight:600;color:var(--text);font-family:'JetBrains Mono',monospace;}
.sidebar-section{background:var(--surface2);border:1px solid var(--border);border-radius:11px;padding:14px;margin-bottom:10px;}
[data-testid="stSelectbox"]>div>div{background:var(--surface2)!important;border-color:var(--border)!important;color:var(--text)!important;}
[data-testid="stExpander"]{background:var(--surface2)!important;border:1px solid var(--border)!important;border-radius:9px!important;}
[data-testid="stExpander"] summary{color:var(--muted)!important;font-size:12px!important;}
[data-testid="stFileUploader"]{background:var(--surface2)!important;border:1px dashed rgba(240,165,0,.3)!important;border-radius:11px!important;}
.welcome-card{border:1px solid var(--border);border-radius:15px;padding:28px;text-align:center;background:var(--surface);margin:32px auto;max-width:540px;}
.welcome-icon{font-size:44px;margin-bottom:14px;}
.welcome-title{font-family:'DM Serif Display',serif;font-size:21px;color:var(--text);margin-bottom:8px;}
.welcome-text{font-size:13.5px;color:var(--muted);line-height:1.6;}
.example-chip{display:inline-block;background:var(--surface2);border:1px solid var(--border);border-radius:18px;padding:5px 13px;font-size:12px;color:var(--muted);margin:3px;}
.upload-hint{background:linear-gradient(135deg,rgba(240,165,0,.06),rgba(224,92,42,.06));border:1px solid rgba(240,165,0,.2);border-radius:10px;padding:12px 16px;font-size:12px;color:#c9a84c;margin-bottom:12px;line-height:1.5;}
.api-hint{background:rgba(240,165,0,.06);border:1px solid rgba(240,165,0,.15);border-radius:9px;padding:11px 14px;font-size:12px;color:#c9a84c;line-height:1.6;margin-top:6px;}
::-webkit-scrollbar{width:5px;}
::-webkit-scrollbar-thumb{background:var(--border);border-radius:3px;}
hr{border-color:var(--border)!important;}
#MainMenu,footer,header{visibility:hidden;}
[data-testid="stToolbar"]{display:none;}
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT PARSERS
# ═══════════════════════════════════════════════════════════════════════════════

def _table_to_markdown(table: list) -> str:
    """Convert a pdfplumber table (list of rows) to a markdown grid."""
    if not table:
        return ""
    rows = []
    for i, row in enumerate(table):
        cells = [str(c).strip() if c else "" for c in row]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("|" + "|".join(["---"] * len(cells)) + "|")
    return "\n".join(rows)


def _docx_table_to_markdown(table) -> str:
    """Convert a python-docx Table object to a markdown grid."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("|" + "|".join(["---"] * len(cells)) + "|")
    return "\n".join(rows)


def _ocr_page(pil_image) -> str:
    """Run Tesseract OCR on a PIL image; return extracted text."""
    try:
        import pytesseract
        return pytesseract.image_to_string(pil_image, config="--psm 6")
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return ""


def parse_pdf(filename: str, file_bytes: bytes) -> tuple[list, dict]:
    """
    Parse a PDF with pdfplumber.
    - Extracts regular text per page
    - Extracts tables as markdown grids
    - Falls back to OCR (Tesseract) for pages with no extractable text

    Returns (documents, parse_stats)
    """
    import pdfplumber
    from pdf2image import convert_from_bytes
    from langchain_core.documents import Document

    docs       = []
    stats      = {"pages": 0, "tables": 0, "ocr_pages": 0, "text_pages": 0}
    ocr_needed = []  # page indices that returned no text

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        stats["pages"] = len(pdf.pages)

        for page_num, page in enumerate(pdf.pages):
            parts = []

            # ── Regular text ──────────────────────────────────────────────────
            raw_text = page.extract_text() or ""

            # ── Tables — extract BEFORE text so we can remove their bbox ──────
            tables = page.extract_tables()
            table_md_blocks = []
            for tbl in tables:
                md = _table_to_markdown(tbl)
                if md:
                    table_md_blocks.append(md)
                    stats["tables"] += 1

            if raw_text.strip():
                parts.append(raw_text.strip())
                stats["text_pages"] += 1
            else:
                ocr_needed.append(page_num)

            for md in table_md_blocks:
                parts.append(f"\n[TABLE]\n{md}\n[/TABLE]\n")

            if parts:
                content = "\n\n".join(parts)
                docs.append(Document(
                    page_content=content,
                    metadata={"source_file": filename, "page": page_num},
                ))

    # ── OCR pass for image-only pages ────────────────────────────────────────
    if ocr_needed:
        try:
            all_images = convert_from_bytes(file_bytes, dpi=200)
            for page_num in ocr_needed:
                if page_num < len(all_images):
                    ocr_text = _ocr_page(all_images[page_num])
                    if ocr_text.strip():
                        docs.append(Document(
                            page_content=f"[OCR]\n{ocr_text.strip()}\n[/OCR]",
                            metadata={"source_file": filename, "page": page_num, "ocr": True},
                        ))
                        stats["ocr_pages"] += 1
        except Exception as e:
            logger.warning(f"OCR pass failed for {filename}: {e}")

    # Sort by page number
    docs.sort(key=lambda d: d.metadata.get("page", 0))
    return docs, stats


def parse_docx(filename: str, file_bytes: bytes) -> tuple[list, dict]:
    """
    Parse a DOCX with python-docx.
    - Extracts paragraphs as text
    - Extracts tables as markdown grids (cell-by-cell, structure preserved)

    Returns (documents, parse_stats)
    """
    import docx as _docx
    from langchain_core.documents import Document

    doc   = _docx.Document(io.BytesIO(file_bytes))
    stats = {"tables": 0, "paragraphs": 0}
    parts = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            # Paragraph
            para = _docx.text.paragraph.Paragraph(block, doc)
            text = para.text.strip()
            if text:
                parts.append(text)
                stats["paragraphs"] += 1

        elif tag == "tbl":
            # Table — convert to markdown
            table = _docx.table.Table(block, doc)
            md    = _docx_table_to_markdown(table)
            if md:
                parts.append(f"\n[TABLE]\n{md}\n[/TABLE]\n")
                stats["tables"] += 1

    content = "\n\n".join(parts)
    documents = [Document(page_content=content, metadata={"source_file": filename})] if content.strip() else []
    return documents, stats


def parse_txt(filename: str, file_bytes: bytes) -> tuple[list, dict]:
    from langchain_core.documents import Document
    text = file_bytes.decode("utf-8", errors="ignore")
    return [Document(page_content=text, metadata={"source_file": filename})], {}


def parse_doc_legacy(filename: str, file_bytes: bytes) -> tuple[list, dict]:
    """Fallback for legacy .doc using docx2txt (no table support)."""
    import docx2txt
    from langchain_core.documents import Document
    text = docx2txt.process(io.BytesIO(file_bytes))
    return [Document(page_content=text, metadata={"source_file": filename})], {"note": "legacy .doc — tables not extracted"}


# ═══════════════════════════════════════════════════════════════════════════════
#  VECTORSTORE
# ═══════════════════════════════════════════════════════════════════════════════

@st.cache_resource(show_spinner=False)
def build_vectorstore_from_bytes(_file_tuples: tuple, cache_version: int):
    """
    cache_version (no leading underscore) is included in Streamlit's cache key —
    incrementing it forces a full rebuild even for identical file bytes.
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import Chroma
    from langchain_community.embeddings import HuggingFaceEmbeddings

    all_docs    = []
    all_stats   = {}

    for filename, file_bytes in _file_tuples:
        ext = Path(filename).suffix.lower()
        try:
            if ext == ".pdf":
                docs, stats = parse_pdf(filename, file_bytes)
            elif ext == ".docx":
                docs, stats = parse_docx(filename, file_bytes)
            elif ext == ".doc":
                docs, stats = parse_doc_legacy(filename, file_bytes)
            elif ext == ".txt":
                docs, stats = parse_txt(filename, file_bytes)
            else:
                logger.warning(f"Unsupported file type: {filename}")
                continue
            all_docs.extend(docs)
            all_stats[filename] = stats
            logger.info(f"Parsed {filename}: {len(docs)} docs, stats={stats}")
        except Exception as e:
            logger.warning(f"Skipped {filename}: {e}")
            all_stats[filename] = {"error": str(e)}

    if not all_docs:
        return None, 0, {}

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, chunk_overlap=200,
        separators=["\n\n", "\n", ". ", "! ", "? "],
    )
    chunks = splitter.split_documents(all_docs)

    embeddings = HuggingFaceEmbeddings(
        model_name=EMBED_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )
    vectorstore = Chroma.from_documents(documents=chunks, embedding=embeddings)
    return vectorstore, len(chunks), all_stats


# ═══════════════════════════════════════════════════════════════════════════════
#  RAG HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def get_mistral_key() -> str:
    try:
        return st.secrets["MISTRAL_API_KEY"]
    except Exception:
        return os.getenv("MISTRAL_API_KEY", "")


def format_sources(source_docs: list) -> str:
    seen, tags = set(), []
    for doc in source_docs:
        name  = doc.metadata.get("source_file", doc.metadata.get("source", "Unknown"))
        page  = doc.metadata.get("page", "")
        ocr   = " 🔍OCR" if doc.metadata.get("ocr") else ""
        label = name + (f" · p{page+1}" if page != "" else "") + ocr
        if label not in seen:
            seen.add(label)
            tags.append(f'<span class="source-tag">📄 {label}</span>')
    return "".join(tags)


def format_docs_labeled(docs: list) -> str:
    parts = []
    for doc in docs:
        source = doc.metadata.get("source_file", "Unknown")
        page   = doc.metadata.get("page", "")
        ocr    = " [OCR]" if doc.metadata.get("ocr") else ""
        label  = source + (f" page {page+1}" if page != "" else "") + ocr
        parts.append(f"[Source: {label}]\n{doc.page_content.strip()}")
    return "\n\n---\n\n".join(parts)


def estimate_confidence(answer: str, source_docs: list) -> str:
    low_phrases = [
        "not found", "not mention", "cannot find", "no information",
        "not available", "does not contain", "not in the document",
        "i don't", "i do not",
    ]
    if any(p in answer.lower() for p in low_phrases):
        return "low"
    distinct = len({d.metadata.get("source_file", "") for d in source_docs})
    if distinct >= 2 and len(answer) > 300:
        return "high"
    return "medium"


def confidence_badge(level: str) -> str:
    labels = {"high": "● High confidence", "medium": "◐ Medium confidence", "low": "○ Low confidence"}
    css    = {"high": "conf-high", "medium": "conf-med", "low": "conf-low"}
    return f'<span class="{css.get(level, "conf-med")}">{labels.get(level, "")}</span>'


def get_cross_doc_context(vectorstore, question: str, indexed_files: list, top_k: int) -> list:
    retriever = vectorstore.as_retriever(
        search_type="mmr",
        search_kwargs={"k": top_k, "fetch_k": top_k * 4, "lambda_mult": 0.65},
    )
    docs = retriever.invoke(question)
    retrieved = {d.metadata.get("source_file", "") for d in docs}
    for source in indexed_files:
        if source not in retrieved:
            try:
                extras = vectorstore.similarity_search(
                    question, k=2, filter={"source_file": source}
                )
                docs.extend(extras)
                retrieved.add(source)
            except Exception:
                pass
    return docs


def build_memory_text(messages: list) -> str:
    qa_pairs = [
        (messages[i], messages[i + 1])
        for i in range(len(messages) - 1)
        if messages[i]["role"] == "user" and messages[i + 1]["role"] == "assistant"
    ]
    recent = qa_pairs[-MEMORY_TURNS:]
    if not recent:
        return ""
    lines = ["Conversation so far (for context):"]
    for u, a in recent:
        lines.append(f"Q: {u['content']}")
        lines.append(f"A: {a['content'][:400]}...")
    return "\n".join(lines)


def export_chat(messages: list) -> str:
    lines = ["SafetyRAG — Chat Export", "=" * 50, ""]
    for msg in messages:
        role = "You" if msg["role"] == "user" else "Assistant"
        lines.append(f"[{role}]")
        lines.append(msg["content"])
        if msg.get("time"):
            lines.append(f"  ⏱ {msg['time']}")
        lines.append("")
    return "\n".join(lines)


def generate_follow_ups(answer: str, question: str, api_key: str, model: str) -> list:
    try:
        from langchain_mistralai import ChatMistralAI
        llm   = ChatMistralAI(model=model, api_key=api_key, temperature=0.4, max_tokens=120)
        prompt = (
            f"Given this Q&A about safety standards:\n"
            f"Q: {question}\nA: {answer[:400]}\n\n"
            f"Suggest exactly 3 short follow-up questions a safety professional might ask next. "
            f"Reply ONLY with the 3 questions, one per line, no numbering or bullets."
        )
        resp  = llm.invoke(prompt)
        lines = [l.strip().lstrip("•-– ") for l in resp.content.strip().split("\n") if l.strip()]
        return lines[:3]
    except Exception:
        return []


def run_rag(
    vectorstore, question: str, model_name: str, api_key: str,
    temperature: float, top_k: int, answer_mode: str,
    indexed_files: list, messages: list, scope_file: str = None,
) -> dict:
    from langchain_core.prompts import PromptTemplate
    from langchain_core.output_parsers import StrOutputParser
    from langchain_mistralai import ChatMistralAI

    llm = ChatMistralAI(
        model=model_name, api_key=api_key,
        temperature=temperature, max_tokens=1536,
    )

    if scope_file and scope_file != "All Documents":
        source_docs = vectorstore.similarity_search(
            question, k=top_k, filter={"source_file": scope_file}
        )
    else:
        source_docs = get_cross_doc_context(vectorstore, question, indexed_files, top_k)

    context     = format_docs_labeled(source_docs)
    memory_text = build_memory_text(messages)
    mode_instr  = ANSWER_MODES.get(answer_mode, ANSWER_MODES["Detailed"])

    template = """\
You are a safety standards expert assistant.
Use ONLY the context below (from official safety standard documents) to answer.
Each chunk is labeled [Source: filename]. Reference the source filename for every fact you state.
The context may include tables in markdown grid format — read them carefully as structured data.
It may also include OCR-extracted text from scanned pages — treat it as authoritative.
If the question spans multiple documents, synthesize information from ALL of them.
If the answer is not in any document, say so clearly — never invent information.

{memory}

Context:
{context}

Question: {question}

Instruction: {mode}

Answer:\
"""
    chain  = PromptTemplate(
        input_variables=["context", "question", "memory", "mode"],
        template=template,
    ) | llm | StrOutputParser()

    answer = chain.invoke({
        "context": context, "question": question,
        "memory": memory_text, "mode": mode_instr,
    })
    return {"result": answer, "source_documents": source_docs}


# ═══════════════════════════════════════════════════════════════════════════════
#  PARSE STATS BADGE RENDERER
# ═══════════════════════════════════════════════════════════════════════════════

def render_parse_stats(parse_stats: dict) -> str:
    """Render a compact per-file parse summary as HTML badges."""
    if not parse_stats:
        return ""
    lines = []
    for fname, stats in parse_stats.items():
        badges = []
        if "error" in stats:
            badges.append(f'<span class="parse-badge" style="color:var(--danger)">⚠ error</span>')
        else:
            if stats.get("pages"):
                badges.append(f'<span class="parse-badge">{stats["pages"]} pages</span>')
            if stats.get("tables"):
                badges.append(f'<span class="parse-badge">📊 {stats["tables"]} tables</span>')
            if stats.get("ocr_pages"):
                badges.append(f'<span class="parse-badge">🔍 {stats["ocr_pages"]} OCR pages</span>')
            if stats.get("paragraphs"):
                badges.append(f'<span class="parse-badge">{stats["paragraphs"]} paragraphs</span>')
            if stats.get("note"):
                badges.append(f'<span class="parse-badge">ℹ {stats["note"]}</span>')
        short = fname if len(fname) < 28 else fname[:25] + "…"
        lines.append(f'<div style="margin-bottom:4px"><span style="font-size:11px;color:var(--muted)">{short}</span> {"".join(badges)}</div>')
    return "".join(lines)


# ═══════════════════════════════════════════════════════════════════════════════
#  SESSION STATE
# ═══════════════════════════════════════════════════════════════════════════════

for key, default in [
    ("messages",      []),
    ("vectorstore",   None),
    ("chunk_count",   0),
    ("indexed_files", []),
    ("index_version", 0),
    ("parse_stats",   {}),
    ("pending_input", ""),
]:
    if key not in st.session_state:
        st.session_state[key] = default


# ═══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════════

with st.sidebar:
    st.markdown("""
    <div style="padding:18px 0 10px">
      <div style="font-family:'DM Serif Display',serif;font-size:20px;color:#e8eaf0">🛡️ SafetyRAG</div>
      <div style="font-size:11px;color:#6b7280;letter-spacing:.08em;text-transform:uppercase;margin-top:3px">
        Cloud Configuration
      </div>
    </div>""", unsafe_allow_html=True)
    st.divider()

    # ── API Key ──
    st.markdown('<span class="sidebar-label">Mistral AI API Key</span>', unsafe_allow_html=True)
    stored_key = get_mistral_key()
    if stored_key:
        st.markdown('<div style="font-size:12px;color:#22c55e;margin-bottom:8px">✅ Key loaded from secrets</div>',
                    unsafe_allow_html=True)
        mistral_key = stored_key
    else:
        mistral_key = st.text_input(
            "Mistral Key", type="password",
            placeholder="Enter your Mistral AI API key...",
            label_visibility="collapsed",
        )
        st.markdown("""
        <div class="api-hint">
          🔑 Get a free key at<br>
          <a href="https://console.mistral.ai" target="_blank" style="color:#f0a500;font-weight:600">
            console.mistral.ai
          </a>
        </div>""", unsafe_allow_html=True)

    st.divider()

    # ── Model & tuning ──
    st.markdown('<span class="sidebar-label">Model</span>', unsafe_allow_html=True)
    model_label = st.selectbox("Model", list(MISTRAL_MODELS.keys()), index=0,
                               label_visibility="collapsed")
    model_name  = MISTRAL_MODELS[model_label]
    temperature = st.slider("Temperature", 0.0, 1.0, 0.1, 0.05, help="Lower = more factual")
    top_k       = st.slider("Top-K Sources", 2, 12, 6, help="Chunks retrieved per question")

    st.divider()

    # ── Answer style ──
    st.markdown('<span class="sidebar-label">Answer Style</span>', unsafe_allow_html=True)
    answer_mode = st.radio("Answer style", list(ANSWER_MODES.keys()), index=1,
                           label_visibility="collapsed", horizontal=True)

    st.divider()

    # ── Upload ──
    st.markdown('<span class="sidebar-label">Upload Documents</span>', unsafe_allow_html=True)
    st.markdown("""
    <div class="upload-hint">
      📁 PDF · DOCX · TXT supported<br>
      Tables, scanned pages &amp; OCR all handled automatically.
    </div>""", unsafe_allow_html=True)

    uploaded_files = st.file_uploader(
        "docs", accept_multiple_files=True,
        type=["pdf", "docx", "doc", "txt"],
        label_visibility="collapsed",
    )

    if uploaded_files:
        if st.button("⚙️ Build Index from Uploads", use_container_width=True):
            with st.spinner(f"Processing {len(uploaded_files)} file(s) — OCR may take a moment..."):
                st.session_state.index_version += 1
                file_tuples = tuple((f.name, f.read()) for f in uploaded_files)
                vs, count, pstats = build_vectorstore_from_bytes(
                    file_tuples, st.session_state.index_version
                )
                if vs:
                    st.session_state.vectorstore   = vs
                    st.session_state.chunk_count   = count
                    st.session_state.indexed_files = [f.name for f in uploaded_files]
                    st.session_state.parse_stats   = pstats
                    st.success(f"✅ {count:,} chunks indexed across {len(uploaded_files)} file(s)!")
                else:
                    st.error("Could not process any files. Check they are valid PDF / DOCX / TXT.")

    st.divider()

    # ── Parse stats ──
    if st.session_state.parse_stats:
        st.markdown('<span class="sidebar-label">Parse Report</span>', unsafe_allow_html=True)
        st.markdown(
            f'<div class="parse-info">{render_parse_stats(st.session_state.parse_stats)}</div>',
            unsafe_allow_html=True,
        )
        st.markdown("")

    # ── Doc scope selector ──
    scope_file = "All Documents"
    if st.session_state.vectorstore and len(st.session_state.indexed_files) > 1:
        st.markdown('<span class="sidebar-label">Search Scope</span>', unsafe_allow_html=True)
        scope_file = st.selectbox(
            "Scope", ["All Documents"] + st.session_state.indexed_files,
            index=0, label_visibility="collapsed",
        )

    # ── Index stats ──
    if st.session_state.vectorstore:
        st.markdown(f"""
        <div class="sidebar-section">
          <div class="sidebar-label">Index Stats</div>
          <div class="stat-row"><span class="stat-key">Chunks</span><span class="stat-val">{st.session_state.chunk_count:,}</span></div>
          <div class="stat-row"><span class="stat-key">Files</span><span class="stat-val">{len(st.session_state.indexed_files)}</span></div>
          <div class="stat-row"><span class="stat-key">Model</span><span class="stat-val">{model_name}</span></div>
          <div class="stat-row"><span class="stat-key">Top-K</span><span class="stat-val">{top_k}</span></div>
          <div class="stat-row"><span class="stat-key">Style</span><span class="stat-val">{answer_mode}</span></div>
        </div>""", unsafe_allow_html=True)

        with st.expander("📄 Indexed files"):
            for fname in st.session_state.indexed_files:
                st.markdown(f'<span class="source-tag">{fname}</span>', unsafe_allow_html=True)

    st.divider()

    # ── Actions ──
    col_a, col_b = st.columns(2)
    with col_a:
        if st.button("🗑️ Clear All", use_container_width=True):
            for k in ["messages", "indexed_files", "parse_stats", "pending_input"]:
                st.session_state[k] = [] if k != "pending_input" and k != "parse_stats" else ({} if k == "parse_stats" else "")
            st.session_state.vectorstore = None
            st.session_state.chunk_count = 0
            st.rerun()
    with col_b:
        if st.session_state.messages:
            st.download_button(
                "💾 Export Chat",
                data=export_chat(st.session_state.messages),
                file_name="safetyrag_chat.txt",
                mime="text/plain",
                use_container_width=True,
            )


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════════

is_ready    = bool(st.session_state.vectorstore) and bool(mistral_key)
scope_label = scope_file if scope_file != "All Documents" else "All Documents"
status_html = (
    f'<span class="status-badge status-ready">● Ready · {scope_label}</span>'
    if is_ready else
    '<span class="status-badge status-loading">● Upload documents</span>'
)

st.markdown(f"""
<div class="rag-header">
  <div class="rag-logo">🛡️</div>
  <div>
    <div class="rag-title">Safety Standards Assistant</div>
    <div class="rag-subtitle">Mistral AI · RAG · ChromaDB · Tables · OCR · Multi-Doc</div>
  </div>
  <div style="margin-left:auto">{status_html}</div>
</div>""", unsafe_allow_html=True)

# ── Welcome ───────────────────────────────────────────────────────────────────
if not st.session_state.messages:
    st.markdown("""
    <div class="welcome-card">
      <div class="welcome-icon">📋</div>
      <div class="welcome-title">Ask me about your safety documents</div>
      <div class="welcome-text">
        Upload PDFs, Word docs, or text files — including scanned documents.<br><br>
        <strong style="color:#f0a500">Tables</strong> are read as structured grids &nbsp;·&nbsp;
        <strong style="color:#f0a500">Scanned pages</strong> are OCR'd automatically &nbsp;·&nbsp;
        <strong style="color:#f0a500">All docs</strong> are synthesized together.
      </div>
      <div style="margin-top:16px">
        <span class="example-chip">PPE requirements for chemical handling</span>
        <span class="example-chip">What does the inspection table show?</span>
        <span class="example-chip">Emergency evacuation steps</span>
        <span class="example-chip">Compare requirements across documents</span>
      </div>
    </div>""", unsafe_allow_html=True)

# ── Chat history ──────────────────────────────────────────────────────────────
else:
    st.markdown('<div class="chat-wrapper">', unsafe_allow_html=True)
    for i, msg in enumerate(st.session_state.messages):
        if msg["role"] == "user":
            st.markdown(f"""
            <div class="msg-user">
              <div class="msg-avatar avatar-user">👤</div>
              <div class="msg-bubble bubble-user">{msg["content"]}</div>
            </div>""", unsafe_allow_html=True)
        else:
            time_html = f'<span class="meta-time">⏱ {msg["time"]}</span>' if msg.get("time") else ""
            conf_html = confidence_badge(msg["confidence"]) if msg.get("confidence") else ""
            srcs_html = (
                f'<div class="sources-block"><strong>Sources</strong><br>{msg["sources"]}</div>'
                if msg.get("sources") else ""
            )
            st.markdown(f"""
            <div class="msg-bot">
              <div class="msg-avatar avatar-bot">🛡️</div>
              <div style="max-width:78%">
                <div class="msg-bubble bubble-bot">{msg["content"]}</div>
                <div class="msg-meta">{time_html}{conf_html}</div>
                {srcs_html}
              </div>
            </div>""", unsafe_allow_html=True)

            # Raw chunk viewer
            if msg.get("chunks"):
                with st.expander(f"🔍 View {len(msg['chunks'])} retrieved chunks"):
                    for j, chunk in enumerate(msg["chunks"]):
                        pg  = f" · p{chunk['page']+1}" if chunk.get("page", "") != "" else ""
                        ocr = " 🔍OCR" if chunk.get("ocr") else ""
                        st.markdown(
                            f'<div class="chunk-label">Chunk {j+1} — {chunk["source"]}{pg}{ocr}</div>',
                            unsafe_allow_html=True,
                        )
                        st.markdown(
                            f'<div class="chunk-box">{chunk["text"][:700]}</div>',
                            unsafe_allow_html=True,
                        )

            # Follow-up chips
            if msg.get("followups"):
                st.markdown("**💡 Suggested follow-ups:**")
                cols = st.columns(len(msg["followups"]))
                for k, fu in enumerate(msg["followups"]):
                    with cols[k]:
                        if st.button(fu, key=f"fu_{i}_{k}", use_container_width=True):
                            st.session_state.pending_input = fu
                            st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)

# ── Input ─────────────────────────────────────────────────────────────────────
st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)

prefill = st.session_state.pending_input
if prefill:
    st.session_state.pending_input = ""

col_input, col_send = st.columns([6, 1])
with col_input:
    user_input = st.text_input(
        "q", label_visibility="collapsed",
        placeholder="e.g. What does the PPE inspection table require for height work?",
        value=prefill, key="user_input", disabled=not is_ready,
    )
with col_send:
    send = st.button("Send →", use_container_width=True, disabled=not is_ready)

# ── Submit ────────────────────────────────────────────────────────────────────
if send and user_input.strip() and is_ready:
    question = user_input.strip()
    st.session_state.messages.append({"role": "user", "content": question})

    with st.spinner("Searching documents and generating answer..."):
        t0     = time.time()
        result = run_rag(
            vectorstore   = st.session_state.vectorstore,
            question      = question,
            model_name    = model_name,
            api_key       = mistral_key,
            temperature   = temperature,
            top_k         = top_k,
            answer_mode   = answer_mode,
            indexed_files = st.session_state.indexed_files,
            messages      = st.session_state.messages,
            scope_file    = scope_file,
        )
        elapsed = time.time() - t0

    answer      = result["result"].strip()
    source_docs = result.get("source_documents", [])

    chunk_data = [
        {
            "source": d.metadata.get("source_file", "Unknown"),
            "page":   d.metadata.get("page", ""),
            "ocr":    d.metadata.get("ocr", False),
            "text":   d.page_content,
        }
        for d in source_docs
    ]

    followups = generate_follow_ups(answer, question, mistral_key, model_name)

    st.session_state.messages.append({
        "role":       "assistant",
        "content":    answer,
        "sources":    format_sources(source_docs),
        "time":       f"{elapsed:.1f}s",
        "confidence": estimate_confidence(answer, source_docs),
        "chunks":     chunk_data,
        "followups":  followups,
    })
    st.rerun()

elif (send or user_input) and not st.session_state.vectorstore:
    st.warning("⚠️ Please upload documents and click **Build Index** first.")
